"""
Generates tailored PDF resumes for each job listing.
Uses python-docx to copy the base DOCX, swap the summary, export via reportlab.
Mirrors the exact logic used locally — only summary changes per job.
"""
import os, re, shutil
from docx import Document

SRC_DOCX = "Jagriti_Mahajan_Resume_1Page.docx"
OUT_DIR  = "resumes"
os.makedirs(OUT_DIR, exist_ok=True)

BASE_SUMMARY = (
    "Senior Technical Business Analyst specializing in financial data platforms, "
    "API ecosystems, and enterprise modernization — 12+ years spanning LSEG, Refinitiv, and Aon. "
    "Drives AI adoption on Snowflake and AWS by translating complex business problems into scalable, "
    "governed data solutions that accelerate delivery and improve enterprise decision-making."
)

# Read current HTML to extract job titles and companies for summary generation
def extract_jobs_from_html():
    try:
        with open("Jagriti_Job_Listings.html","r") as f:
            html = f.read()
        titles   = re.findall(r'class="job-title">(.*?)</div>', html)
        companies= re.findall(r'class="company">(.*?)</div>', html)
        return list(zip(titles[:12], companies[:12]))
    except:
        return []

def make_summary(title, company):
    title   = re.sub(r'<[^>]+>','', title).strip()
    company = re.sub(r'<[^>]+>','', company).strip()

    # Generic tailored summary based on job title keywords
    t = title.lower()
    if "financial" in t or "finance" in t:
        return (f"Senior Technical Business Analyst with 12+ years in financial data platforms, "
                f"enterprise analytics, and stakeholder-driven decision support — spanning LSEG, Refinitiv, and Aon. "
                f"Expert in SQL, Power BI, and data governance on Snowflake and AWS, with a track record of "
                f"translating financial requirements into scalable, governed data solutions.")
    elif "data" in t or "analytics" in t:
        return (f"Senior Technical Business Analyst with 12+ years delivering data platforms, "
                f"API ecosystems, and analytics solutions in enterprise environments. "
                f"Deep expertise in SQL, Power BI, Snowflake, and AWS with a strong record of translating "
                f"complex business needs into actionable data insights for {company}.")
    elif "api" in t or "integration" in t:
        return (f"Senior Technical Business Analyst with 12+ years specializing in API design, "
                f"system integration, and enterprise modernization — spanning LSEG, Refinitiv, and Aon. "
                f"Expert in RESTful APIs, JSON schemas, Swagger, and end-to-end data flows on Snowflake and AWS.")
    elif "agile" in t or "scrum" in t or "product" in t:
        return (f"Senior Technical Business Analyst with 12+ years driving Agile delivery, "
                f"backlog management, and cross-functional team alignment across enterprise programs. "
                f"Proven record owning sprint planning for 35-member teams and translating epics into "
                f"well-defined user stories with clear acceptance criteria at LSEG, Refinitiv, and Aon.")
    else:
        return BASE_SUMMARY

def build_pdf(docx_path, pdf_path):
    """Export DOCX to PDF using reportlab by extracting text from DOCX."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_RIGHT

    DARK_BLUE = colors.HexColor('#0f3d78')
    MED_BLUE  = colors.HexColor('#1a56a4')
    MUTED     = colors.HexColor('#4b5563')
    BLACK     = colors.HexColor('#1f2937')

    def S(name, **kw):
        defaults = dict(fontName='Helvetica', fontSize=9.5, textColor=BLACK,
                        spaceAfter=1.5, leading=12.5)
        defaults.update(kw)
        return ParagraphStyle(name, **defaults)

    styles = {
        'name':    S('name', fontName='Helvetica-Bold', fontSize=15, textColor=DARK_BLUE, spaceAfter=2, leading=18),
        'contact': S('contact', fontSize=8.5, textColor=MUTED, spaceAfter=4, leading=11),
        'sec':     S('sec', fontName='Helvetica-Bold', fontSize=9, textColor=MED_BLUE, spaceBefore=5, spaceAfter=2, leading=11),
        'body':    S('body', fontSize=9, spaceAfter=2, leading=12),
        'bullet':  S('bullet', fontSize=8.8, spaceAfter=1.5, leading=11.5, leftIndent=12, firstLineIndent=-8),
        'note':    S('note', fontName='Helvetica-Oblique', fontSize=8, textColor=MUTED, spaceAfter=2, leading=10),
        'edu':     S('edu', fontSize=8.8, spaceAfter=1, leading=11),
    }

    def hr(story):
        story.append(HRFlowable(width='100%', thickness=0.8,
                                color=colors.HexColor('#dbeafe'), spaceAfter=2, spaceBefore=0))

    def sec_title(story, t):
        story.append(Spacer(1, 2))
        story.append(Paragraph(t, styles['sec']))
        hr(story)

    def jrow(story, title, co, date, client=None):
        left = f'<b>{title}</b><br/><font color="#4b5563" size="8">{co}</font>'
        if client:
            left += f'<br/><i><font color="#4b5563" size="7.5">{client}</font></i>'
        tbl = Table([[Paragraph(left, styles['body']),
                      Paragraph(f'<b>{date}</b>', styles['note'])]],
                    colWidths=[4.8*inch, 2.2*inch])
        tbl.setStyle(TableStyle([
            ('VALIGN',(0,0),(-1,-1),'TOP'),('ALIGN',(1,0),(1,0),'RIGHT'),
            ('LEFTPADDING',(0,0),(-1,-1),0),('RIGHTPADDING',(0,0),(-1,-1),0),
            ('TOPPADDING',(0,0),(-1,-1),2),('BOTTOMPADDING',(0,0),(-1,-1),1),
        ]))
        story.append(tbl)

    def bul(story, text):
        story.append(Paragraph(f'• {text}', styles['bullet']))

    def erow(story, deg, school, date, note=None):
        tbl = Table([[Paragraph(f'<b>{deg}</b> — <font color="#4b5563">{school}</font>', styles['edu']),
                      Paragraph(f'<font color="#4b5563">{date}</font>', styles['edu'])]],
                    colWidths=[5.3*inch, 1.8*inch])
        tbl.setStyle(TableStyle([
            ('VALIGN',(0,0),(-1,-1),'TOP'),('ALIGN',(1,0),(1,0),'RIGHT'),
            ('LEFTPADDING',(0,0),(-1,-1),0),('RIGHTPADDING',(0,0),(-1,-1),0),
            ('TOPPADDING',(0,0),(-1,-1),2),('BOTTOMPADDING',(0,0),(-1,-1),1),
        ]))
        story.append(tbl)
        if note:
            story.append(Paragraph(note, styles['note']))

    # Read all text from the DOCX
    doc = Document(docx_path)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    summary = paras[3] if len(paras) > 3 else BASE_SUMMARY

    def get_bullets(start_marker, end_marker):
        capturing, bullets = False, []
        for p in paras:
            if start_marker.lower() in p.lower(): capturing = True; continue
            if capturing and end_marker.lower() in p.lower(): break
            if capturing and p: bullets.append(p)
        return bullets

    betanxt_bullets = get_bullets("BETANXT", "ADROIT")
    adroit_bullets  = get_bullets("ADROIT",  "Aon Hewitt")
    aon_bullets     = get_bullets("Aon Hewitt", "EDUCATION")

    doc_out = SimpleDocTemplate(pdf_path, pagesize=letter,
        topMargin=0.45*inch, bottomMargin=0.45*inch,
        leftMargin=0.55*inch, rightMargin=0.55*inch)
    story = []

    story.append(Paragraph("Jagriti Mahajan", styles['name']))
    story.append(Paragraph(
        "Brookfield, WI  |  (908) 240-9209  |  jagriti.mahajan0507@gmail.com  |  linkedin.com/in/jagritimahajan0507",
        styles['contact']))

    sec_title(story, "PROFESSIONAL SUMMARY")
    story.append(Paragraph(summary, styles['body']))

    sec_title(story, "TECHNICAL SKILLS")
    skills = [
        ("Cloud & Data Platforms",       "AWS, Snowflake, Confluent Cloud"),
        ("API & Integration",            "RESTful APIs, JSON, Swagger, ReadyAPI, Swagger Editor"),
        ("Data & Analytics",             "SQL, Power BI, Tableau, Informatica, SAS Enterprise Guide, MS Excel, Data.World"),
        ("Methodologies",                "Agile, Scrum, SAFe, UAT, BRD/FRD"),
        ("Database & Operating Systems", "IBM DB2, RDBMS, SQL Server, DynamoDB, MySQL, Mainframe, Datacom, Windows"),
        ("Modeling & Designing",         "MS Visio, SmartDraw, Process Model, LucidChart"),
        ("Other Tools",                  "IBM Audit Viewer, JIRA, Confluence, GitLab, Maestro"),
    ]
    rows = []
    for i in range(4):
        lk, lv = skills[i]
        right = f'<b>{skills[i+4][0]}:</b> {skills[i+4][1]}' if i+4 < len(skills) else ''
        rows.append([Paragraph(f'<b>{lk}:</b> {lv}', styles['body']),
                     Paragraph(right, styles['body'])])
    stbl = Table(rows, colWidths=[3.55*inch, 3.55*inch])
    stbl.setStyle(TableStyle([
        ('VALIGN',(0,0),(-1,-1),'TOP'),
        ('LEFTPADDING',(0,0),(-1,-1),0),('RIGHTPADDING',(0,0),(-1,-1),4),
        ('TOPPADDING',(0,0),(-1,-1),0),('BOTTOMPADDING',(0,0),(-1,-1),2),
    ]))
    story.append(stbl)

    sec_title(story, "WORK EXPERIENCE")
    jrow(story, "Sr. Technical Business Analyst", "BETANXT (formerly LSEG) — Brookfield, WI", "Oct 2021 – Present")
    for b in betanxt_bullets: bul(story, b)
    story.append(Spacer(1, 3))
    jrow(story, "Sr. Technical Business Analyst", "ADROIT — Brookfield, WI", "May 2018 – Oct 2021", client="Client: Refinitiv")
    for b in adroit_bullets: bul(story, b)
    story.append(Spacer(1, 3))
    jrow(story, "Business Systems Analyst", "Aon Hewitt", "Nov 2011 – Jul 2016")
    for b in aon_bullets: bul(story, b)

    sec_title(story, "EDUCATION")
    erow(story, "MS in Computer Information Systems", "Colorado State University, USA", "Dec 2017",
         note="Graduate Web Developer (Part-Time), Aug 2016 – Dec 2017")
    erow(story, "Bachelor of Technology in Information Technology",
         "Lovely Professional University, India", "May 2011")

    doc_out.build(story)

def replace_summary_in_docx(docx_path, new_summary):
    doc = Document(docx_path)
    para = doc.paragraphs[3]
    if para.runs:
        first = para.runs[0]
        sz, nm = first.font.size, first.font.name
        for r in para.runs: r.text = ''
        first.text = new_summary
        first.bold = False
        if sz: first.font.size = sz
        if nm: first.font.name = nm
    doc.save(docx_path)

if __name__ == "__main__":
    import json

    if not os.path.exists(SRC_DOCX):
        print(f"Base DOCX not found: {SRC_DOCX}")
        exit(1)

    jobs = extract_jobs_from_html()
    if not jobs:
        print("No jobs found in HTML — generating 12 generic resumes.")
        jobs = [("Senior Business Analyst", "Company") for _ in range(12)]

    names = [
        "resume_job_01","resume_job_02","resume_job_03","resume_job_04",
        "resume_job_05","resume_job_06","resume_job_07","resume_job_08",
        "resume_job_09","resume_job_10","resume_job_11","resume_job_12",
    ]

    changes = {}

    for i, (title, company) in enumerate(jobs[:12]):
        name     = names[i]
        tmp_docx = os.path.join(OUT_DIR, name + ".docx")
        pdf_path = os.path.join(OUT_DIR, name + ".pdf")
        shutil.copy2(SRC_DOCX, tmp_docx)
        new_summary = make_summary(title, company)
        replace_summary_in_docx(tmp_docx, new_summary)
        build_pdf(tmp_docx, pdf_path)
        os.remove(tmp_docx)

        changed = new_summary.strip() != BASE_SUMMARY.strip()
        changes[name] = {
            "title":        re.sub(r'<[^>]+>', '', title).strip(),
            "company":      re.sub(r'<[^>]+>', '', company).strip(),
            "changed":      changed,
            "base_summary": BASE_SUMMARY,
            "new_summary":  new_summary,
        }
        status = "updated" if changed else "unchanged"
        print(f"  ✓ {name}.pdf  [{status}]  ({title[:50]})")

    with open(os.path.join(OUT_DIR, "changes.json"), "w") as f:
        json.dump(changes, f, indent=2)
    print("All resumes generated. Changes written to resumes/changes.json.")
